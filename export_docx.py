#!/usr/bin/env python3
"""Export a markdown report to a formatted Word (.docx) document.

Usage:
    python export_docx.py reports/prospectus_decatur_commerce_park.md
    python export_docx.py reports/prospectus_decatur_commerce_park.md -o ~/Desktop/report.docx
"""

import argparse
import os
import re
import sys

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap


def _set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)


def _style_table(table):
    """Apply professional styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if i == 0:
                _set_cell_shading(cell, '1F3864')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True


def md_to_docx(md_path, output_path=None):
    """Convert a markdown report to a Word document."""
    with open(md_path) as f:
        lines = f.readlines()

    if not output_path:
        output_path = md_path.replace('.md', '.docx')

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Heading styles
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(31, 56, 100)
        if level == 1:
            hs.font.size = Pt(20)
        elif level == 2:
            hs.font.size = Pt(14)
        else:
            hs.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin line
            run = p.add_run('_' * 80)
            run.font.color.rgb = RGBColor(200, 200, 200)
            run.font.size = Pt(6)
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Table
        if line.startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # Filter out separator rows
            data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|$', l)]

            if len(data_lines) < 1:
                continue

            # Parse cells
            rows_data = []
            for tl in data_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                rows_data.append(cells)

            if not rows_data:
                continue

            ncols = len(rows_data[0])
            table = doc.add_table(rows=len(rows_data), cols=ncols)
            table.style = 'Table Grid'

            for ri, row_data in enumerate(rows_data):
                for ci, cell_text in enumerate(row_data):
                    if ci < ncols:
                        cell = table.rows[ri].cells[ci]
                        # Clean markdown formatting
                        clean = cell_text.replace('**', '').replace('`', '')
                        cell.text = clean

            _style_table(table)
            doc.add_paragraph()  # spacing after table
            continue

        # Bullet list
        if line.startswith('- '):
            text = line[2:].strip()
            # Handle checkbox items
            text = text.replace('[ ]', '☐').replace('[x]', '☑')
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, text)
            i += 1
            continue

        # Indented sub-bullet
        if line.startswith('  - '):
            text = line[4:].strip()
            p = doc.add_paragraph(style='List Bullet 2')
            _add_formatted_runs(p, text)
            i += 1
            continue

        # Italic block (starts with _)
        if line.startswith('_') and line.endswith('_'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('_'))
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            i += 1
            continue

        # Regular paragraph with formatting
        p = doc.add_paragraph()
        _add_formatted_runs(p, line)
        i += 1

    doc.save(output_path)
    return output_path


def _add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = paragraph._element.makeelement(qn('w:hyperlink'), {qn('r:id'): r_id})
    new_run = paragraph._element.makeelement(qn('w:r'), {})
    rPr = paragraph._element.makeelement(qn('w:rPr'), {})
    rStyle = paragraph._element.makeelement(qn('w:rStyle'), {qn('w:val'): 'Hyperlink'})
    color = paragraph._element.makeelement(qn('w:color'), {qn('w:val'): '1F3864'})
    underline = paragraph._element.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    sz = paragraph._element.makeelement(qn('w:sz'), {qn('w:val'): '20'})
    rPr.append(rStyle)
    rPr.append(color)
    rPr.append(underline)
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def _add_formatted_runs(paragraph, text):
    """Parse markdown bold/italic/code/links and add as formatted runs."""
    # Split on **bold**, *italic*, `code`, [link](url) patterns
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif part.startswith('[') and '](' in part:
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                _add_hyperlink(paragraph, m.group(1), m.group(2))
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def main():
    parser = argparse.ArgumentParser(description="Export markdown report to Word docx")
    parser.add_argument("input", help="Input markdown file")
    parser.add_argument("-o", "--output", help="Output docx path (default: same name, .docx extension)")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: {args.input} not found")
        sys.exit(1)

    output = args.output or args.input.replace('.md', '.docx')
    result = md_to_docx(args.input, output)
    print(f"Word document saved to: {result}")


if __name__ == "__main__":
    main()
